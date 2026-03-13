#!/usr/bin/env python3
"""
Convert exam.md to a two-column Word document with checkboxes.

Usage:
  exam/.venv/bin/python exam_to_docx.py
  exam/.venv/bin/python exam_to_docx.py --input exam.md --output exam.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")  # ~0.25 inch column gap


def _set_compact_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for style_name in ("Normal", "Heading 1", "Heading 2"):
        s = doc.styles[style_name]
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.0


def _is_question_line(line: str) -> bool:
    stripped = line.strip()
    return stripped[:1].isdigit() and ". " in stripped


def _is_option_line(line: str) -> bool:
    stripped = line.strip()
    return len(stripped) > 2 and stripped[1:3] == ". " and stripped[0] in "ABCD"


def _add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def convert_to_docx(md_text: str, output_path: Path) -> None:
    doc = Document()
    _set_compact_layout(doc)
    _set_two_columns(doc.sections[0])

    in_code = False

    for line in md_text.splitlines():
        if line.strip() == "Answer Key":
            break

        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            _add_code_paragraph(doc, line.rstrip("\n"))
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue

        if _is_question_line(line):
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(2)
            continue

        if _is_option_line(line):
            p = doc.add_paragraph(f"☐ {line.strip()}")
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)
            continue

        if line.strip() == "":
            continue

        doc.add_paragraph(line.strip())

    doc.save(str(output_path))


def main() -> int:
    script_dir = Path(__file__).resolve().parent
    parser = argparse.ArgumentParser(description="Convert exam.md to Word.")
    parser.add_argument("--input", default=str(script_dir / "exam.md"))
    parser.add_argument("--output", default=str(script_dir / "exam.docx"))
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}")

    md_text = input_path.read_text(encoding="utf-8")
    convert_to_docx(md_text, output_path)
    print(f"Wrote {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
